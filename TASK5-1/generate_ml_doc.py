"""生成机器学习算法与评价指标 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ===== 全局样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    hf = heading_style.font
    hf.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        hf.size = Pt(22)
    elif level == 2:
        hf.size = Pt(16)
    else:
        hf.size = Pt(13)

# ===== 辅助函数 =====
def add_paragraph(doc, text, bold=False, italic=False, size=None, color=None, alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = size
    if color: run.font.color.rgb = color
    if alignment: p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ===== 封面 =====
doc.add_paragraph()
doc.add_paragraph()
add_paragraph(doc, '机器学习分类算法与评价指标', bold=True, size=Pt(28),
              color=RGBColor(0x1A, 0x1A, 0x2E), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '—— 逻辑回归 · 决策树 · 随机森林 · XGBoost 等', size=Pt(14),
              color=RGBColor(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '', size=Pt(6))
add_paragraph(doc, '混淆矩阵 · ROC曲线 · AUC · 精确率 · 召回率 · F1-Score', size=Pt(14),
              color=RGBColor(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_paragraph(doc, '2025年7月', size=Pt(12), color=RGBColor(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ===== 目录页 =====
doc.add_heading('目录', level=1)
toc_items = [
    '第一部分：分类机器学习算法',
    '    1.1 逻辑回归 (Logistic Regression)',
    '    1.2 决策树 (Decision Tree)',
    '    1.3 随机森林 (Random Forest)',
    '    1.4 XGBoost',
    '    1.5 LightGBM',
    '    1.6 支持向量机 (SVM)',
    '    1.7 K近邻 (KNN)',
    '    1.8 多层感知机 (MLP)',
    '    1.9 算法选择指南',
    '第二部分：机器学习模型评价指标',
    '    2.1 混淆矩阵',
    '    2.2 准确率、精确率、召回率、F1-Score',
    '    2.3 ROC 曲线与 AUC',
    '    2.4 各类指标的使用场景',
    '    2.5 模型评估的最佳实践',
    '附录：结合 TASK5-1 数据集的实践示例',
]
for item in toc_items:
    size = Pt(14) if not item.startswith('    ') else Pt(11)
    bold = not item.startswith('    ')
    color = RGBColor(0x1A, 0x1A, 0x2E) if bold else RGBColor(0x55, 0x55, 0x55)
    add_paragraph(doc, item, bold=bold, size=size, color=color, space_after=Pt(3))
doc.add_page_break()

# ======================================================================
# 第一部分：分类机器学习算法
# ======================================================================
doc.add_heading('第一部分：分类机器学习算法', level=1)

add_paragraph(doc, '分类是机器学习中最基础、最常见的任务之一。分类算法的目标是根据输入特征，将样本划分到预定义的离散类别中。例如：判断一封邮件是否为垃圾邮件（二分类）、识别手写数字 0-9（多分类）、预测股票次日涨跌（二分类）。')
add_paragraph(doc, '分类算法按照原理大致可分为以下几类：', bold=True)
add_paragraph(doc, '• 线性模型：逻辑回归、线性判别分析\n• 树模型：决策树、随机森林、XGBoost、LightGBM\n• 距离模型：K近邻（KNN）、支持向量机（SVM）\n• 神经网络：多层感知机（MLP）')

# --- 1.1 逻辑回归 ---
doc.add_heading('1.1 逻辑回归 (Logistic Regression)', level=2)

add_paragraph(doc, '核心思想', bold=True, size=Pt(12))
add_paragraph(doc, '逻辑回归虽然名字里有"回归"，但实际上是一个分类算法。它的核心思想是：用线性模型计算一个分数（log-odds），然后通过 Sigmoid 函数将其映射到 [0, 1] 区间，作为分类概率。')

add_paragraph(doc, '数学公式：')
add_paragraph(doc, 'P(Y=1|X) = 1 / (1 + e^(-z))', italic=True, size=Pt(11))
add_paragraph(doc, '其中 z = w₀ + w₁x₁ + w₂x₂ + ... + wₙxₙ，即特征的线性组合。')

add_paragraph(doc, '以 TASK5-1 数据集为例：', italic=True)
add_paragraph(doc, '假设我们要预测财报公布次日股价是否上涨（Y=True/False），逻辑回归会为每个特征分配一个系数：')
add_paragraph(doc, '• 营业总收入增速 → 系数 +0.153（正向，增速越高越容易涨）\n• 市净率 PB → 系数 -0.129（负向，PB 越高越容易跌）\n• 基本每股收益增速 → 系数 +0.081（正向）')

add_paragraph(doc, '优点与局限', bold=True, size=Pt(12))
add_styled_table(doc, ['优点', '局限'], [
    ['可解释性极强：每个系数的正负号和大小直接反映特征对结果的影响方向与程度', '只能捕捉线性关系，无法自动发现非线性模式'],
    ['输出概率值，便于设定不同的分类阈值', '对共线性敏感，需要提前处理高相关特征'],
    ['训练速度快，内存占用小，适合大规模数据', '特征工程依赖人工（如标准化、交互项构造）'],
    ['支持 L1/L2 正则化，可自动做特征选择', '在复杂问题上的精度通常不如集成树模型'],
])

# --- 1.2 决策树 ---
doc.add_heading('1.2 决策树 (Decision Tree)', level=2)

add_paragraph(doc, '核心思想', bold=True, size=Pt(12))
add_paragraph(doc, '决策树通过"反复提问"的方式做分类。从根节点开始，每次选择一个最优特征和一个分割点，将数据分成两个子集，递归下去，直到满足停止条件。整个过程就像一棵倒置的树。')

add_paragraph(doc, '决策树的分裂过程：')
add_paragraph(doc, '┌─────────────┐\n│ 全部样本    │ ← 根节点\n│ Y=False: 60% │\n│ Y=True:  40% │\n└──────┬──────┘\n       │ 股息率 > 2%？\n  ┌────┴────┐\n  ▼         ▼\n┌─────┐  ┌─────┐\n│ NO  │  │ YES │  ← 内部节点\n│ F多  │  │ T多  │\n└──┬──┘  └──┬──┘\n   │净利润  │EPS\n   │增速>0? │增速>10%?\n  ┌┴┐    ┌┴┐\n ┌┐┌┐   ┌┐┌┐   ← 叶节点（最终分类）', size=Pt(9))

add_paragraph(doc, '常用的分裂准则：', bold=True)
add_paragraph(doc, '• 信息增益（ID3）：选择使信息熵下降最多的特征\n• 信息增益率（C4.5）：在信息增益基础上惩罚分支过多的特征\n• 基尼系数（CART）：衡量节点"不纯度"，值越小越纯')

add_paragraph(doc, '优点与局限', bold=True, size=Pt(12))
add_styled_table(doc, ['优点', '局限'], [
    ['完全可解释：可以通过规则路径追溯每个预测的推理过程', '容易过拟合：树太深时会记住训练数据的噪声'],
    ['无需特征标准化，可同时处理数值型和类别型特征', '不稳定：数据的微小变化可能导致完全不同的树结构'],
    ['自动捕捉非线性关系和特征交互', '倾向于偏向取值多的特征（可用信息增益率缓解）'],
    ['计算开销小', '单棵树精度通常不如集成方法'],
])

# --- 1.3 随机森林 ---
doc.add_heading('1.3 随机森林 (Random Forest)', level=2)

add_paragraph(doc, '核心思想', bold=True, size=Pt(12))
add_paragraph(doc, '随机森林是"集成学习"的代表算法。它的哲学是：一百个普通人投票比一个专家更可靠。具体做法是训练很多棵决策树（通常 100~500 棵），每棵树用不同的数据子集和特征子集训练，最终分类结果由所有树投票决定。')

add_paragraph(doc, '两个"随机"：', bold=True)
add_paragraph(doc, '1. 样本随机（Bootstrap）：每棵树从原始数据中随机有放回地抽取同样数量的样本。大约 36.8% 的样本不会被抽中（称为 OOB，袋外样本），可用作天然的验证集。')
add_paragraph(doc, '2. 特征随机：每次分裂时，只从随机选取的一部分特征（通常 sqrt(总特征数) 个）中选择最优分裂点。这迫使不同的树关注不同的特征组合。')

add_paragraph(doc, '在 TASK5-1 中，随机森林是 AUC 最高的模型（0.625）：', italic=True)
add_paragraph(doc, '• 训练了 200 棵树，每棵最大深度 10\n• 17 个全量特征，特征集 A（16 特征）表现几乎相同（AUC 差 0.002）\n• 对共线性不敏感，无需 Winsorize 之外的额外处理')

add_paragraph(doc, '优点与局限', bold=True, size=Pt(12))
add_styled_table(doc, ['优点', '局限'], [
    ['精度高、鲁棒性强，通常优于单棵决策树', '模型较大，训练和推理速度比单棵树慢'],
    ['可输出特征重要性，辅助特征筛选', '决策边界不够平滑，无法外推'],
    ['OOB 误差可替代交叉验证，省去额外验证集', '对高维稀疏数据表现不如 GBDT 类模型'],
    ['不易过拟合：Bootstrap + 随机特征 = 天然的方差削减', '黑盒程度高于单棵树，SHAP 等工具辅助解释'],
])

# --- 1.4 XGBoost ---
doc.add_heading('1.4 XGBoost', level=2)

add_paragraph(doc, '核心思想', bold=True, size=Pt(12))
add_paragraph(doc, 'XGBoost（eXtreme Gradient Boosting）属于"提升法"（Boosting）。与随机森林的"并行投票"不同，Boosting 是"串行接力"：每一棵新树专门去修正前面所有树犯的错误（拟合残差），逐步提升整体精度。')

add_paragraph(doc, '关键创新点：')
add_paragraph(doc, '• 正则化：在目标函数中显式加入树的复杂度惩罚项（叶节点数 + 权重平方和），防止过拟合\n• 二阶泰勒展开：比传统 GBDT 只用一阶梯度信息更精确\n• 列抽样（Column Subsampling）：借鉴随机森林，每次分裂只考虑部分特征\n• 早停（Early Stopping）：验证集性能不再提升时自动停止训练，避免过拟合\n• 缺失值自动处理：自动学习最优的缺失值分裂方向')

add_paragraph(doc, '在 TASK5-1 中：', italic=True)
add_paragraph(doc, '• XGBoost (B) AUC = 0.6195，略低于随机森林\n• 使用了 scale_pos_weight 处理类别不平衡\n• Recall = 0.539（比随机森林高），更擅长找出"会涨"的样本\n• LightGBM 性能接近（AUC = 0.614），但训练更快')

add_paragraph(doc, '优点与局限', bold=True, size=Pt(12))
add_styled_table(doc, ['优点', '局限'], [
    ['Kaggle 竞赛常胜算法，精度通常最高', '超参数多，调优需要经验和时间'],
    ['内置正则化，过拟合风险低于传统 GBDT', '对异常值较敏感（有 L1/L2 正则化和 Huber Loss 缓解）'],
    ['支持 GPU 加速、分布式训练', '小数据集上可能不如简单模型'],
    ['可处理缺失值，无需预处理填充', '模型解释需借助 SHAP（比 LR 复杂得多）'],
])

# --- 1.5 LightGBM ---
doc.add_heading('1.5 LightGBM', level=2)

add_paragraph(doc, 'LightGBM 是微软开源的梯度提升框架，与 XGBoost 同属 Boosting 家族。它的核心区别在于树的生长策略：')
add_paragraph(doc, '• XGBoost：按层生长（Level-wise），每层所有节点同时分裂\n• LightGBM：按叶生长（Leaf-wise），每次选择增益最大的叶子分裂')
add_paragraph(doc, 'Leaf-wise 策略使得 LightGBM 收敛更快、精度更高，但也更可能过拟合（需要控制 max_depth 和 num_leaves）。')

add_paragraph(doc, 'LightGBM 的独特优势：', bold=True)
add_paragraph(doc, '• GOSS（Gradient-based One-Side Sampling）：保留大梯度样本，随机采样小梯度样本，加速训练\n• EFB（Exclusive Feature Bundling）：将互斥的稀疏特征捆绑，减少特征维度\n• 原生支持类别特征，无需 One-Hot 编码')

# --- 1.6 SVM ---
doc.add_heading('1.6 支持向量机 (SVM)', level=2)

add_paragraph(doc, '核心思想', bold=True, size=Pt(12))
add_paragraph(doc, 'SVM 的目标是找到一个"最大间隔"超平面，将两类样本尽可能分开，并且离超平面最近的样本点（支持向量）到平面的距离最大化。对于线性不可分的问题，通过"核函数"将数据映射到高维空间，在高维中找到线性分割面。')

add_paragraph(doc, '常用核函数：')
add_paragraph(doc, '• 线性核（Linear）：就是原始空间的线性分类器\n• 多项式核（Polynomial）：映射到多项式特征空间\n• RBF 核（Radial Basis Function）：最常用，映射到无限维空间\n• Sigmoid 核：类似神经网络的激活函数')

add_paragraph(doc, '优点与局限', bold=True, size=Pt(12))
add_styled_table(doc, ['优点', '局限'], [
    ['在高维数据上表现优秀（如文本分类）', '大规模数据训练很慢（O(n²) ~ O(n³)）'],
    ['核技巧可处理非线性问题', '需要仔细的特征标准化'],
    ['只依赖支持向量，对噪声有一定鲁棒性', '不直接输出概率（需 Platt Scaling 校准）'],
    ['理论基础扎实，有泛化误差界', '核函数和参数选择依赖经验'],
])

# --- 1.7 KNN ---
doc.add_heading('1.7 K近邻 (K-Nearest Neighbors)', level=2)

add_paragraph(doc, 'KNN 是最简单的分类算法：对于一个新样本，找到训练集中离它最近的 K 个邻居，看这些邻居中哪一类的占比最高，就预测为该类。')

add_paragraph(doc, '关键参数：')
add_paragraph(doc, '• K 值：太小容易受噪声影响，太大边界模糊（通常用交叉验证选择）\n• 距离度量：欧氏距离（最常用）、曼哈顿距离、余弦相似度')

add_paragraph(doc, '优点与局限', bold=True, size=Pt(12))
add_styled_table(doc, ['优点', '局限'], [
    ['原理极其简单，无需训练（惰性学习）', '预测时需要计算所有训练样本的距离，速度慢'],
    ['天然支持多分类', '对特征量纲敏感，必须标准化'],
    ['无显式假设，只要数据够多就能逼近任意决策边界', '高维数据下距离度量失效（维度灾难）'],
])

# --- 1.8 MLP ---
doc.add_heading('1.8 多层感知机 (MLP)', level=2)

add_paragraph(doc, 'MLP 是最基础的前馈神经网络，由输入层、若干隐藏层、输出层组成。每层神经元通过加权连接，经激活函数（ReLU、tanh、sigmoid）引入非线性。')

add_paragraph(doc, '激活函数：')
add_paragraph(doc, '• ReLU(x) = max(0, x)：最常用，缓解梯度消失\n• Sigmoid(x) = 1/(1+e^(-x))：将值压缩到 (0,1)，用于二分类输出层\n• Softmax：多分类输出层，将 logits 转为概率分布')

add_paragraph(doc, '在 TASK5-1 中，MLP AUC = 0.580，在 2 万条数据上表现一般——神经网络通常需要更大的数据量。')

# --- 1.9 算法选择指南 ---
doc.add_heading('1.9 算法选择指南', level=2)

add_styled_table(doc, ['场景', '推荐算法', '原因'], [
    ['需要解释每个特征的边际效应', '逻辑回归', '系数符号和大小直接可读'],
    ['精度最重要，可解释性次要', 'XGBoost / LightGBM', 'Kaggle 竞赛级性能'],
    ['精度与鲁棒性兼顾', '随机森林', '不易过拟合，OOB 自带验证'],
    ['数据量小（< 1 万）、特征少', '逻辑回归 / SVM', '简单模型在小数据上更稳定'],
    ['数据量大（> 10 万）、特征多', 'XGBoost / LightGBM / MLP', '复杂模型优势更明显'],
    ['需要快速原型验证', 'LightGBM', '训练速度和精度俱佳'],
    ['特征存在大量共线性', '随机森林 / XGBoost', '树模型对共线性不敏感'],
    ['重视可解释性的金融/医疗场景', '逻辑回归 / 单棵决策树', '合规和解释要求严格'],
    ['高维稀疏数据（如文本）', 'SVM + 线性核 / 逻辑回归 + L1', '稀疏性天然适合线性模型'],
])

doc.add_page_break()

# ======================================================================
# 第二部分：机器学习模型评价指标
# ======================================================================
doc.add_heading('第二部分：机器学习模型评价指标', level=1)

add_paragraph(doc, '选择合适的评价指标和选择合适的模型同样重要。不同的业务场景需要不同的指标来评判模型的优劣。')

# --- 2.1 混淆矩阵 ---
doc.add_heading('2.1 混淆矩阵 (Confusion Matrix)', level=2)

add_paragraph(doc, '混淆矩阵是分类模型最基础的评估工具，以表格形式展示预测值与真实值的交叉统计。', bold=True)

add_paragraph(doc, '二分类混淆矩阵结构：')
add_styled_table(doc, ['', '预测为 Positive', '预测为 Negative'], [
    ['实际为 Positive', 'TP (True Positive)\n真正例 · 正确识别', 'FN (False Negative)\n假负例 · 漏报'],
    ['实际为 Negative', 'FP (False Positive)\n假正例 · 误报', 'TN (True Negative)\n真负例 · 正确排除'],
])

add_paragraph(doc, '以 TASK5-1 为例（XGBoost 在测试集上的表现）：', italic=True)
add_paragraph(doc, '• TP = 434：模型预测"涨"，实际涨了（抓到了 53.9% 的上涨机会）\n• FP = 1310：模型预测"涨"，实际跌了（误判率偏高）\n• FN = 372：模型预测"跌"，实际涨了（漏掉了 46.1% 的上涨）\n• TN = 2146：模型预测"跌"，实际跌了（正确排除了大部分下跌）')

add_paragraph(doc, '从混淆矩阵可以派生出所有其他指标：')
add_paragraph(doc, '• 准确率 Accuracy = (TP+TN) / (TP+TN+FP+FN)\n• 精确率 Precision = TP / (TP+FP)\n• 召回率 Recall = TP / (TP+FN)\n• 特异度 Specificity = TN / (TN+FP)\n• F1-Score = 2 × Precision × Recall / (Precision + Recall)')

# --- 2.2 四大核心指标 ---
doc.add_heading('2.2 准确率 · 精确率 · 召回率 · F1-Score', level=2)

add_paragraph(doc, '准确率 (Accuracy)', bold=True, size=Pt(12))
add_paragraph(doc, '定义：预测正确的样本占总样本的比例。')
add_paragraph(doc, 'Accuracy = (TP + TN) / (TP + TN + FP + FN)')
add_paragraph(doc, '⚠️ 陷阱：在类别不平衡时，高准确率可能具有欺骗性。例如在 TASK5-1 中，如果模型永远预测"不涨"（False），准确率也有 81%，但 AUC = 0.50——模型什么也没学到。')

add_paragraph(doc, '精确率 (Precision)', bold=True, size=Pt(12))
add_paragraph(doc, '定义：预测为"涨"的样本中，实际涨了的比例。')
add_paragraph(doc, 'Precision = TP / (TP + FP)')
add_paragraph(doc, '适用场景：当你关心"如果我说涨，有多大把握真的涨"。例如，选股策略中如果只有 10 个仓位，你要确保买入的股票确实会涨。')

add_paragraph(doc, '召回率 (Recall)', bold=True, size=Pt(12))
add_paragraph(doc, '定义：实际涨了的样本中，被模型正确识别出来的比例。')
add_paragraph(doc, 'Recall = TP / (TP + FN)')
add_paragraph(doc, '适用场景：当你关心"有多少上涨机会被我抓住了"。例如，量化基金希望尽可能多地捕获涨幅，即使意味着会有一些误判。')

add_paragraph(doc, 'F1-Score', bold=True, size=Pt(12))
add_paragraph(doc, '定义：精确率和召回率的调和平均数。')
add_paragraph(doc, 'F1 = 2 × Precision × Recall / (Precision + Recall)')
add_paragraph(doc, '适用场景：当你需要同时平衡 Precision 和 Recall，且两类错误代价差不多时。调和平均确保任一指标过低都会显著拉低 F1。')

add_paragraph(doc, 'Precision vs Recall 的权衡（Trade-off）：', bold=True)
add_paragraph(doc, '模型输出的概率需要设定一个阈值（默认 0.5）来决定分类：\n• 提高阈值 → Precision↑, Recall↓（只选最有把握的"涨"，但会漏掉很多真上涨）\n• 降低阈值 → Recall↑, Precision↓（尽可能抓住上涨，但误判增多）\n具体选多少阈值取决于业务需求。')

# --- 2.3 ROC 与 AUC ---
doc.add_heading('2.3 ROC 曲线与 AUC', level=2)

add_paragraph(doc, 'ROC 曲线 (Receiver Operating Characteristic)', bold=True, size=Pt(12))
add_paragraph(doc, 'ROC 曲线是一张展示模型在所有可能分类阈值下表现的图：')
add_paragraph(doc, '• X 轴：假阳性率 FPR = FP / (FP + TN) — 即"实际跌却被判涨"的比例\n• Y 轴：真阳性率 TPR = TP / (TP + FN) — 即"实际涨并被识别"的比例（也就是 Recall）')
add_paragraph(doc, '曲线上每一个点对应一个分类阈值。阈值从 0 变化到 1，点从右上角移到左下角，画出一条曲线。')

add_paragraph(doc, '从 TASK5-1 的 ROC 图可以看到：', italic=True)
add_paragraph(doc, '• 随机森林（深红色实线）在最左上角，整体最接近理想模型\n• 所有模型的 ROC 都明显高于对角线（随机猜测线）\n• 低 FPR 区域（左侧），各模型 TPR 约在 0.10~0.15，区分度有限')

add_paragraph(doc, 'AUC (Area Under the Curve)', bold=True, size=Pt(12))
add_paragraph(doc, 'AUC 是 ROC 曲线下的面积，取值范围 [0.5, 1.0]：')
add_paragraph(doc, '• AUC = 0.50：模型跟随机猜测一样（对角线）\n• AUC = 0.60 ~ 0.70：有一定的区分能力（TASK5-1 大部分模型在此区间）\n• AUC = 0.70 ~ 0.80：较好的模型\n• AUC = 0.80 ~ 0.90：优秀的模型\n• AUC = 0.90 ~ 1.00：接近完美的模型（需警惕过拟合或数据泄漏）')

add_paragraph(doc, 'AUC 的统计学含义：', bold=True)
add_paragraph(doc, 'AUC 有一个非常直观的解释：随机选一个正样本和一个负样本，模型给正样本的打分高于负样本的概率。AUC = 0.625 意味着有 62.5% 的概率，模型能把"会涨"的样本排在"不会涨"的样本前面。')

add_paragraph(doc, 'AUC vs Accuracy：何时用哪个？')
add_paragraph(doc, '• 类别均衡 + 阈值为 0.5 → Accuracy 即可\n• 类别不均衡（如 TASK5-1：涨:跌 ≈ 4:6）→ AUC 更可靠，因为 AUC 不受阈值影响\n• 需要排序能力而不只是分类 → AUC\n• 需要根据预测概率选 top-N 样本 → AUC（本质上评估的是排序质量）')

add_paragraph(doc, 'PR 曲线（补充指标）：', bold=True)
add_paragraph(doc, '当正负样本极度不平衡时（如欺诈检测，正样本仅 0.1%），ROC 可能过于乐观。此时 PR 曲线（Precision-Recall Curve）更能反映模型真实能力。TASK5-1 的不平衡程度（4:6）尚可接受，ROC 足以胜任。')

# --- 2.4 指标选择场景 ---
doc.add_heading('2.4 各类指标的使用场景', level=2)

add_styled_table(doc, ['业务场景', '首选指标', '说明'], [
    ['股票涨跌预测（TASK5-1）', 'AUC', '关注排序质量，不固定阈值'],
    ['疾病筛查（癌症检测）', 'Recall', '漏诊代价极高，宁可误诊也要找出所有患者'],
    ['垃圾邮件过滤', 'Precision', '误判正常邮件为垃圾比漏判垃圾邮件更糟糕'],
    ['推荐系统 Top-N', 'Precision@N / NDCG', '只关心排名靠前的推荐质量'],
    ['信用评分 / 风控', 'AUC + KS', '需要整体排序能力 + 区分度'],
    ['类别均衡的一般分类', 'Accuracy + F1', '简单直接，所有错误成本相同'],
    ['极度不平衡（欺诈、罕见病）', 'AUPRC (PR 曲线下面积)', 'ROC 在极度不平衡下过于乐观'],
])

# --- 2.5 最佳实践 ---
doc.add_heading('2.5 模型评估的最佳实践', level=2)

add_paragraph(doc, '1. 永远不要在训练集上评估模型', bold=True)
add_paragraph(doc, '训练集指标会显著高估真实性能。必须使用独立的测试集。')

add_paragraph(doc, '2. 时间序列数据不能随机划分', bold=True)
add_paragraph(doc, '股票数据、销售数据等有天然的时间顺序，随机划分会导致"用未来预测过去"的信息泄漏。TASK5-1 采用按日期划分：训练集 2021-06 ~ 2022-03，测试集 2022-04 ~ 2022-06。')

add_paragraph(doc, '3. 交叉验证提供更稳健的评估', bold=True)
add_paragraph(doc, 'K-Fold 交叉验证将数据分成 K 份，轮流用 K-1 份训练、1 份测试。报告 K 次结果的均值和标准差，可评估模型稳定性。时间序列数据应使用 TimeSeriesSplit。')

add_paragraph(doc, '4. 不要只看一个指标', bold=True)
add_paragraph(doc, '例如 Accuracy 高但 AUC 低 → 模型只是记住了多数类。同时关注 AUC、Accuracy、Precision、Recall、F1，全面评估。')

add_paragraph(doc, '5. 设置合理的基线', bold=True)
add_paragraph(doc, '任何模型至少应该超越最简单的基线：\n• 随机猜测基线：AUC = 0.50\n• 多数类基线：Accuracy = 最大类占比\n• 简单规则基线：如"永远预测 PE<20 的股票涨"')

add_paragraph(doc, '6. 分析模型弱点', bold=True)
add_paragraph(doc, '不仅看总体指标，还要分段分析：\n• 不同行业/板块的准确率是否一致？\n• 不同市场环境（牛市/熊市）下表现如何？（稳健性检验）\n• 预测概率最高的 10% 样本准确率是多少？（校准度）')

doc.add_page_break()

# ======================================================================
# 附录
# ======================================================================
doc.add_heading('附录：TASK5-1 数据集建模结果汇总', level=1)

add_paragraph(doc, '数据集概况', bold=True, size=Pt(12))
add_styled_table(doc, ['项目', '说明'], [
    ['样本量', '20,772 条'],
    ['股票数', '4,281 只'],
    ['特征数', '17 个（估值 9 个 + 成长性 8 个）'],
    ['时间范围', '2021-06-30 ~ 2022-06-30'],
    ['目标变量', 'Y: 财报公布次日涨跌（盘后发布，无前视偏差）'],
    ['Train/Test', '按日期划分：2021-06~2022-03 训练 / 2022-04~2022-06 测试'],
    ['特征工程', 'Winsorize 1%/99% → VIF 共线性剔除（仅删利润总额增速）'],
])

add_paragraph(doc, '模型对比结果', bold=True, size=Pt(12))
add_styled_table(doc, ['排名', '模型', '特征集', 'AUC', 'Accuracy', 'F1'], [
    ['🥇', 'Random Forest', 'B (全量)', '0.6251', '69.71%', '0.3403'],
    ['🥈', 'Random Forest', 'A (去共线)', '0.6228', '69.47%', '0.3399'],
    ['🥉', 'XGBoost', 'B (全量)', '0.6195', '60.49%', '0.3406'],
    ['4', 'LightGBM', 'A (去共线)', '0.6140', '67.39%', '0.3272'],
    ['5', 'XGBoost', 'A (去共线)', '0.6137', '61.83%', '0.3536'],
    ['6', 'LightGBM', 'B (全量)', '0.6135', '67.29%', '0.3292'],
    ['7', 'Logistic Regression', 'A (去共线)', '0.5815', '75.20%', '0.2675'],
    ['8', 'MLP', 'A (去共线)', '0.5795', '67.01%', '0.3115'],
    ['-', 'Baseline (Dummy)', '-', '0.5000', '81.07%', '0.0000'],
])

add_paragraph(doc, '关键发现：', bold=True)
add_paragraph(doc, '1. 树模型（RF、XGBoost、LightGBM）整体优于线性模型和神经网络\n2. 共线性剔除（A vs B）对树模型几乎没有影响（AUC 差 < 0.006）\n3. 估值特征提供约 0.03 AUC 的增量信息（纯成长 XGB AUC = 0.584）\n4. LR 系数可解释：营收增速(+0.15)正向、市净率(-0.13)负向\n5. 所有模型的 AUC 在 0.58~0.63 之间，说明基于财报数据的次日涨跌预测有一定信号但难度较大')

# ===== 保存 =====
output_path = r'C:\Users\RYAN\Desktop\BA工作坊\TASK5-1\机器学习分类算法与评价指标.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
print(f'File size: {os.path.getsize(output_path) / 1024:.1f} KB')
